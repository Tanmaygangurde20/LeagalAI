import streamlit as st
import os
import time
import uuid
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import io
from docx import Document

# Load environment variables
load_dotenv()

# Import agent classes
from drafting.graph import LegalDocumentAgent, AgentState
from clarification.graphSearch import LegalSearchGraph
from clarification.summarize import LegalSummarizer
from document_qa.graphRag import DocumentQARAG

# --- UI CONFIG ---
st.set_page_config(page_title="Unified Legal AI Suite", page_icon="⚖️", layout="wide")

# --- SESSION STATE INIT ---
def init_session():
    if 'drafting_agent' not in st.session_state:
        st.session_state.drafting_agent = LegalDocumentAgent()
    if 'drafting_session_id' not in st.session_state:
        st.session_state.drafting_session_id = str(uuid.uuid4())
    if 'drafting_document' not in st.session_state:
        st.session_state.drafting_document = None
    if 'clarification_search' not in st.session_state:
        st.session_state.clarification_search = LegalSearchGraph(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            gemini_api_key=os.getenv("GEMINI_API_KEY")
        )
    if 'clarification_summarizer' not in st.session_state:
        st.session_state.clarification_summarizer = LegalSummarizer(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            gemini_api_key=os.getenv("GEMINI_API_KEY")
        )
    if 'clarification_history' not in st.session_state:
        st.session_state.clarification_history = []
    if 'docqa_rag' not in st.session_state:
        st.session_state.docqa_rag = DocumentQARAG(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            google_api_key=os.getenv("GEMINI_API_KEY")
        )
    if 'docqa_processed' not in st.session_state:
        st.session_state.docqa_processed = False
    if 'docqa_current_document' not in st.session_state:
        st.session_state.docqa_current_document = None
    if 'docqa_chat_history' not in st.session_state:
        st.session_state.docqa_chat_history = []

init_session()

# --- SIDEBAR NAVIGATION ---
pages = [
    "🏠 Home",
    "📝 Document Drafting",
    "🔍 Legal Clarification",
    "📄 Document-Based QA"
]
page = st.sidebar.radio("Navigation", pages, index=0)

st.sidebar.markdown("---")
st.sidebar.markdown("<small style='color:#888;'>Unified Legal AI Suite &copy; 2024</small>", unsafe_allow_html=True)

# --- PAGE 1: HOME ---
if page == "🏠 Home":
    st.markdown('<h1 style="text-align:center; font-size:2.5rem; color:#1f4e79;">⚖️ Unified Legal AI Suite</h1>', unsafe_allow_html=True)
    st.markdown("""
    <div style='text-align:center; color:#444; margin-bottom:2rem;'>
        <b>Draft legal documents, get legal clarifications, and ask questions about your own documents—all in one place.</b>
    </div>
    <div style='max-width:800px; margin:auto; font-size:1.1rem; color:#222; background:#f8f9fa; padding:2rem; border-radius:1rem; box-shadow:0 2px 8px rgba(0,0,0,0.04);'>
        <h3>Project Overview</h3>
        <p>This suite is a production-ready, agentic AI platform for legal professionals and users. It provides three core capabilities:</p>
        <ol>
            <li><b>Conversational Legal Document Drafting:</b> Generate well-structured legal documents (NDA, contracts, leases) by simply describing your needs. The AI collects all required information and produces a downloadable document.</li>
            <li><b>Legal Clarification:</b> Ask law-related questions and receive real-time, source-cited answers. The system searches trusted legal sources, extracts keywords, and summarizes results using advanced LLMs.</li>
            <li><b>Document-Based QA:</b> Upload your own legal documents (PDF, DOCX, TXT) and ask context-specific questions. The system retrieves relevant sections and provides accurate, grounded answers.</li>
        </ol>
        <p>All modules are powered by state-of-the-art LLMs and vector search, with robust fallback and error handling. No legal advice is provided—educational and informational use only.</p>
    </div>
    """, unsafe_allow_html=True)

# --- PAGE 2: DOCUMENT DRAFTING ---
elif page == "📝 Document Drafting":
    st.markdown('<h2 style="color:#1f4e79;">📝 Conversational Legal Document Drafting</h2>', unsafe_allow_html=True)
    st.info("Interactively draft NDAs, contracts, or lease agreements. The AI will ask for missing details and generate a complete document. Your chat history will be shown below.")
    def reset_drafting():
        st.session_state.drafting_session_id = str(uuid.uuid4())
        st.session_state.drafting_document = None
        st.session_state.drafting_chat = []
        if 'drafting_state_dict' in st.session_state:
            del st.session_state['drafting_state_dict']
        if 'drafting_question_index' in st.session_state:
            del st.session_state['drafting_question_index']
    st.button("Start New Drafting", on_click=reset_drafting)
    agent = st.session_state.drafting_agent
    if 'drafting_chat' not in st.session_state:
        st.session_state.drafting_chat = []
    # Only show input if no document is generated yet
    if not st.session_state.drafting_document:
        if 'drafting_state_dict' not in st.session_state:
            user_input = st.text_input("Describe the document you want to draft (e.g., 'Draft an NDA between Alice and Bob'):", key="drafting_input")
            if user_input:
                state = AgentState(session_id=st.session_state.drafting_session_id, user_input=user_input)
                state_dict = state.model_dump()
                state_dict = agent.identify_document_type(state_dict)
                st.session_state.drafting_state_dict = state_dict
                st.session_state.drafting_answers = {}
                st.session_state.drafting_question_index = 0
                st.session_state.drafting_chat.append({"role": "user", "content": user_input})
                st.rerun()
        else:
            state_dict = st.session_state.drafting_state_dict
            # Sequentially ask questions
            while not state_dict.get("is_complete", False):
                state_dict = agent.ask_question(state_dict)
                question = state_dict.get("current_question", "")
                if question:
                    q_idx = st.session_state.get('drafting_question_index', 0)
                    key = f"drafting_q_{st.session_state.drafting_session_id}_{q_idx}"
                    answer = st.text_input(f"{question}", key=key)
                    if answer:
                        state_dict["user_input"] = answer
                        state_dict = agent.process_answer(state_dict)
                        st.session_state.drafting_question_index = q_idx + 1
                        st.session_state.drafting_state_dict = state_dict
                        st.session_state.drafting_chat.append({"role": "ai", "content": question})
                        st.session_state.drafting_chat.append({"role": "user", "content": answer})
                        st.rerun()
                    else:
                        break
            # Generate document if complete
            if state_dict.get("is_complete", False):
                state_dict = agent.generate_document(state_dict)
                document = state_dict.get("final_document", "[No document generated]")
                st.session_state.drafting_document = document
                st.session_state.drafting_chat.append({"role": "ai", "content": "---\n**Generated Legal Document**\n" + document})
                del st.session_state['drafting_state_dict']
                del st.session_state['drafting_question_index']
                st.success("Document generated! See below.")
    # Show chat history with toggle
    if 'show_drafting_chat' not in st.session_state:
        st.session_state.show_drafting_chat = True
    if st.session_state.drafting_chat:
        col1, col2 = st.columns([1, 8])
        with col1:
            if st.session_state.show_drafting_chat:
                if st.button('Hide Chat History', key='hide_drafting_chat'):
                    st.session_state.show_drafting_chat = False
            else:
                if st.button('Show Chat History', key='show_drafting_chat'):
                    st.session_state.show_drafting_chat = True
        with col2:
            if st.session_state.show_drafting_chat:
                st.markdown("---")
                st.subheader("Chat History")
                # Only show Q&A, not the generated document
                for msg in st.session_state.drafting_chat:
                    if msg["role"] == "user":
                        st.markdown(f"**You:** {msg['content']}")
                    elif msg["role"] == "ai" and not msg["content"].startswith("---\n**Generated Legal Document**"):
                        st.markdown(f"**AI:** {msg['content']}")
    # Show the generated document and download button
    if st.session_state.drafting_document:
        st.markdown("---")
        # Format the document for markdown preview
        def format_legal_doc_to_markdown(doc_text):
            lines = doc_text.splitlines()
            md_lines = []
            for line in lines:
                line = line.strip()
                if not line:
                    md_lines.append("")
                    continue
                # Bold all-caps section headers and key phrases
                if line.isupper() and len(line) > 3:
                    md_lines.append(f"**{line.title()}**")
                elif line.startswith("LANDLORD:") or line.startswith("TENANT:") or line.startswith("PROPERTY:") or line.startswith("LEASE TERM:") or line.startswith("MONTHLY RENT:") or line.startswith("SECURITY DEPOSIT:") or line.startswith("USE OF PREMISES:") or line.startswith("CONDITION OF PREMISES:") or line.startswith("REPAIRS AND MAINTENANCE:") or line.startswith("UTILITIES:") or line.startswith("INSPECTIONS:") or line.startswith("DEFAULT:") or line.startswith("NOTICES:") or line.startswith("GOVERNING LAW:") or line.startswith("ENTIRE AGREEMENT:") or line.startswith("AMENDMENTS:") or line.startswith("BINDING EFFECT:") or line.startswith("ACKNOWLEDGMENT:") or line.startswith("SIGNATURES:"):
                    parts = line.split(":", 1)
                    if len(parts) == 2:
                        md_lines.append(f"**{parts[0]}:**{parts[1]}")
                    else:
                        md_lines.append(f"**{line}**")
                else:
                    md_lines.append(line)
            return "\n".join(md_lines)
        formatted_doc = format_legal_doc_to_markdown(st.session_state.drafting_document)
        st.markdown(f"{formatted_doc}")

        # Download as Word (.docx)
        def generate_docx(doc_text):
            doc = Document()
            lines = doc_text.splitlines()
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                # Bold all-caps section headers and key phrases
                if line.isupper() and len(line) > 3:
                    p = doc.add_paragraph()
                    run = p.add_run(line.title())
                    run.bold = True
                elif line.startswith("LANDLORD:") or line.startswith("TENANT:") or line.startswith("PROPERTY:") or line.startswith("LEASE TERM:") or line.startswith("MONTHLY RENT:") or line.startswith("SECURITY DEPOSIT:") or line.startswith("USE OF PREMISES:") or line.startswith("CONDITION OF PREMISES:") or line.startswith("REPAIRS AND MAINTENANCE:") or line.startswith("UTILITIES:") or line.startswith("INSPECTIONS:") or line.startswith("DEFAULT:") or line.startswith("NOTICES:") or line.startswith("GOVERNING LAW:") or line.startswith("ENTIRE AGREEMENT:") or line.startswith("AMENDMENTS:") or line.startswith("BINDING EFFECT:") or line.startswith("ACKNOWLEDGMENT:") or line.startswith("SIGNATURES:"):
                    parts = line.split(":", 1)
                    if len(parts) == 2:
                        p = doc.add_paragraph()
                        run = p.add_run(parts[0] + ":")
                        run.bold = True
                        p.add_run(parts[1])
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                else:
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf
        docx_buf = generate_docx(st.session_state.drafting_document)
        st.download_button(
            label="💾 Download Document (.docx)",
            data=docx_buf,
            file_name=f"legal_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- PAGE 3: LEGAL CLARIFICATION ---
elif page == "🔍 Legal Clarification":
    st.markdown('<h2 style="color:#1f4e79;">🔍 Legal Clarification</h2>', unsafe_allow_html=True)
    st.info("Ask a legal question. See sample questions below, extracted keywords, search results, and the LLM-generated summary.")
    sample_questions = [
        "What is the difference between void and voidable contracts in Canadian law?",
        "What are the requirements for adverse possession in Ontario?",
        "How is child custody determined in Canadian family courts?",
        "What constitutes negligence under Canadian tort law?"
    ]
    st.markdown("**Sample Questions:**")
    for q in sample_questions:
        st.markdown(f"- {q}")
    query = st.text_area(
        "Enter your legal question:",
        value=st.session_state.get("clarification_current_query", ""),
        height=80,
        placeholder="e.g., What is the difference between void and voidable contracts in Canadian law?",
        key="clarification_query_input"
    )
    col1, col2 = st.columns([1, 1])
    with col1:
        search_button = st.button("🔎 Search & Summarize", key="clarification_search_btn")
    with col2:
        clear_button = st.button("🗑️ Clear", key="clarification_clear_btn")
    if clear_button:
        st.session_state.clarification_current_query = ""
        st.rerun()
    if search_button and query.strip():
        st.session_state.clarification_current_query = query
        progress_bar = st.progress(0)
        status_text = st.empty()
        try:
            status_text.text("Searching legal databases...")
            progress_bar.progress(25)
            search_results = st.session_state.clarification_search.search_legal_query(query)
            if not search_results.get("success", False):
                st.error(f"Search failed: {search_results.get('error', 'Unknown error')}")
            else:
                status_text.text("Summarizing results...")
                progress_bar.progress(75)
                summary_results = st.session_state.clarification_summarizer.generate_comprehensive_summary(search_results)
                progress_bar.progress(100)
                status_text.text("Done!")
                time.sleep(0.5)
                progress_bar.empty()
                status_text.empty()
                st.markdown("---")
                st.markdown(f"**Your Question:** {query}")
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown("**Extracted Keywords:**")
                keywords = search_results.get("keywords", [])
                if keywords:
                    st.markdown(", ".join([f"`{kw}`" for kw in keywords]))
                else:
                    st.markdown("_No keywords extracted._")
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown("**Search Results:**")
                results = search_results.get("results", [])
                if results:
                    for i, result in enumerate(results, 1):
                        st.markdown(f"**{i}.** {result.get('content', '')[:300]}{'...' if len(result.get('content', ''))>300 else ''}")
                        st.caption(f"Source: {result.get('source', 'Unknown')}")
                else:
                    st.markdown("_No search results found._")
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown("**LLM-Generated Summary:**")
                if summary_results.get("success", False):
                    st.success(summary_results.get('summary', ''))
                    if summary_results.get("citations"):
                        st.markdown("**Citations:**")
                        for i, citation in enumerate(summary_results["citations"], 1):
                            st.markdown(f"{i}. {citation}")
                else:
                    st.error(f"Summary failed: {summary_results.get('error', 'Unknown error')}")
        except Exception as e:
            st.error(f"Processing failed: {str(e)}")
            progress_bar.empty()
            status_text.empty()

# --- PAGE 4: DOCUMENT-BASED QA ---
elif page == "📄 Document-Based QA":
    st.markdown('<h2 style="color:#1f4e79;">📄 Document-Based QA</h2>', unsafe_allow_html=True)
    st.info("Upload a legal document (PDF, DOCX, or TXT) and ask context-specific questions.")
    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_file = st.file_uploader(
            "Choose a document",
            type=['pdf', 'docx', 'doc', 'txt'],
            help="Upload a PDF, DOCX, or TXT file"
        )
    with col2:
        if uploaded_file is not None:
            st.info(f"**File Details:**\n- Name: {uploaded_file.name}\n- Size: {uploaded_file.size / 1024:.1f} KB\n- Type: {uploaded_file.type}")
    if uploaded_file is not None:
        if st.button("🚀 Process Document", key="docqa_process_btn"):
            with st.spinner("Processing document..."):
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                file_extension = uploaded_file.name.split('.')[-1].lower()
                try:
                    result = st.session_state.docqa_rag.process_document_and_query(
                        file_path=tmp_file_path,
                        file_type=file_extension,
                        query="What is this document about?"
                    )
                except Exception as e:
                    result = {'success': False, 'error': str(e)}
                os.unlink(tmp_file_path)
                if result['success']:
                    st.session_state.docqa_processed = True
                    st.session_state.docqa_current_document = uploaded_file.name
                    st.success(f"Document processed! You can now ask questions about: {uploaded_file.name}")
                    st.write("**Initial Analysis:**")
                    st.write(result['answer'])
                else:
                    st.session_state.docqa_processed = False
                    st.session_state.docqa_current_document = None
                    error_msg = result.get('error', 'Unknown error occurred')
                    if error_msg is None:
                        error_msg = 'Unknown error occurred'
                    if isinstance(error_msg, str):
                        if 'no healthy upstream' in error_msg or 'No LLM could be initialized' in error_msg:
                            st.error("All LLMs are currently unavailable. Please try again later or check your API keys.")
                        else:
                            st.error(f"Error processing document: {error_msg}")
                    else:
                        st.error("An unknown error occurred while processing the document.")
    if st.session_state.docqa_processed:
        st.markdown("---")
        st.subheader("💬 Ask Questions About Your Document")
        question = st.text_input(
            "Ask a question about your document:",
            placeholder="e.g., What are the main obligations of the parties?",
            key="docqa_question_input"
        )
        ask_button = st.button("🤔 Ask Question", key="docqa_ask_btn")
        clear_history = st.button("🗑️ Clear History", key="docqa_clear_btn")
        if clear_history:
            st.session_state.docqa_chat_history = []
            st.rerun()
        if ask_button and question:
            with st.spinner("Thinking..."):
                result = st.session_state.docqa_rag.query_existing_documents(question)
            if result['success']:
                timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
                st.session_state.docqa_chat_history.append((question, result['answer'], timestamp))
                st.success(result['answer'])
            else:
                st.error(f"Error: {result.get('error', 'Unknown error occurred')}")
        if st.session_state.docqa_chat_history:
            st.subheader("💭 Chat History")
            for i, (q, a, timestamp) in enumerate(reversed(st.session_state.docqa_chat_history)):
                with st.expander(f"Q{len(st.session_state.docqa_chat_history)-i}: {q[:50]}..." if len(q) > 50 else f"Q{len(st.session_state.docqa_chat_history)-i}: {q}", expanded=(i==0)):
                    st.write(f"**Question:** {q}")
                    st.write(f"**Answer:** {a}")
                    st.caption(f"Asked at: {timestamp}") 