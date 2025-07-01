import streamlit as st
import os
import time
import uuid
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
import io
from docx import Document

# Load environment variables
load_dotenv()

# Import agent classes
from drafting.graph import LegalDocumentAgent, AgentState
from clarification.graphSearch import LegalSearchGraph
from clarification.summarize import LegalSummarizer
from document_qa.graphRag import DocumentQARAG

# --- ADVANCED UI STYLES ---
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    html, body, [class*="css"]  {
        font-family: 'Inter', sans-serif !important;
        background: linear-gradient(135deg, ##e080d4 0%, #e3e9f7 100%) !important;
    }
    .main-header {
        font-size: 2.8rem;
        color: #f5f7f2;
        font-weight: 700;
        text-align: center;
        margin-bottom: 0.5rem;
        letter-spacing: 1px;
        animation: fadeInDown 1s;
    }
    .suite-logo {
        width: 60px;
        margin: 0 auto 1rem auto;
        display: block;
        animation: fadeIn 1.2s;
    }
    .section-card {
        background: linear-gradient(90deg, #1976d2 0%, #1a237e 100%);
        border-radius: 1.2rem;
        box-shadow: 0 4px 24px 0 rgba(26,35,126,0.08);
        padding: 2rem 2.5rem;
        margin-bottom: 2rem;
        animation: fadeInUp 0.8s;
    }
    .stButton > button {
        background: linear-gradient(90deg, #1a237e 0%, #1976d2 100%);
        color: #fff;
        border: none;
        border-radius: 0.5rem;
        font-weight: 600;
        font-size: 1.1rem;
        padding: 0.7rem 1.5rem;
        box-shadow: 0 2px 8px 0 rgba(26,35,126,0.08);
        transition: background 0.2s, transform 0.2s;
    }
    .stButton > button:hover {
        background: linear-gradient(90deg, #1976d2 0%, #1a237e 100%);
        transform: translateY(-2px) scale(1.03);
    }
    .stTextInput > div > div > input {
        border-radius: 0.5rem;
        border: 1.5px solid #e3e9f7;
        font-size: 1.1rem;
        padding: 0.6rem 1rem;
        background: ##e080d4;
    }
    .stTextArea textarea {
        border-radius: 0.5rem;
        border: 1.5px solid #e3e9f7;
        font-size: 1.1rem;
        background: ##e080d4;
    }
    .stDownloadButton > button {
        background: linear-gradient(90deg, #ffd600 0%, #ffb300 100%);
        color: #1a237e;
        font-weight: 700;
        border-radius: 0.5rem;
        font-size: 1.1rem;
        margin-top: 1rem;
        box-shadow: 0 2px 8px 0 rgba(255,214,0,0.08);
        transition: background 0.2s, transform 0.2s;
    }
    .stDownloadButton > button:hover {
        background: linear-gradient(90deg, #ffb300 0%, #ffd600 100%);
        transform: translateY(-2px) scale(1.03);
    }
    .stAlert, .stInfo, .stSuccess, .stError {
        border-radius: 0.7rem !important;
        font-size: 1.08rem !important;
    }
    .chat-bubble {
        background: #e3e9f7;
        border-radius: 1rem;
        padding: 1rem 1.5rem;
        margin-bottom: 0.7rem;
        box-shadow: 0 2px 8px 0 rgba(26,35,126,0.04);
        animation: fadeInUp 0.7s;
    }
    .chat-bubble.user {
        background: #fffde7;
        color: #1a237e;
        border-left: 4px solid #ffd600;
    }
    .chat-bubble.ai {
        background: #e3e9f7;
        color: #1a237e;
        border-left: 4px solid #1976d2;
    }
    .divider {
        border: none;
        border-top: 2px solid #e3e9f7;
        margin: 2rem 0 1.5rem 0;
    }
    .footer {
        text-align: center;
        color: #888;
        font-size: 1rem;
        margin-top: 2.5rem;
        padding-bottom: 1rem;
        letter-spacing: 0.5px;
    }
    @keyframes fadeInDown {
        0% { opacity: 0; transform: translateY(-30px); }
        100% { opacity: 1; transform: translateY(0); }
    }
    @keyframes fadeInUp {
        0% { opacity: 0; transform: translateY(30px); }
        100% { opacity: 1; transform: translateY(0); }
    }
    @keyframes fadeIn {
        0% { opacity: 0; }
        100% { opacity: 1; }
    }
    /* Provided sidebar and radio button styling */
    .css-1d391kg {
        background: linear-gradient(180deg, #2c3e50 0%, #34495e 100%);
        border-radius: 0 20px 20px 0;
    }
    .css-17eq0hr {
        background: linear-gradient(180deg, #2c3e50 0%, #34495e 100%);
    }
    .sidebar .sidebar-content {
        background: transparent;
    }
    .stRadio > div {
        gap: 0.5rem;
    }
    .stRadio > div > label {
        background: rgba(255, 255, 255, 0.1);
        padding: 12px 16px;
        border-radius: 12px;
        color: #ecf0f1;
        font-weight: 500;
        transition: all 0.3s ease;
        cursor: pointer;
        border: 2px solid transparent;
        backdrop-filter: blur(10px);
    }
    .stRadio > div > label:hover {
        background: rgba(255, 255, 255, 0.2);
        transform: translateX(5px);
        border-color: rgba(255, 255, 255, 0.3);
    }
    .stRadio > div > label[data-checked="true"] {
        background: linear-gradient(135deg, #3498db, #2980b9);
        transform: translateX(8px);
        box-shadow: 0 4px 15px rgba(52, 152, 219, 0.4);
    }
    .feature-card {
        background: linear-gradient(90deg, #1976d2 0%, #1a237e 100%);
        border-radius: 18px;
        padding: 2rem;
        margin: 1.5rem 0;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
        border: 2.5px solid #fff;
        backdrop-filter: blur(10px);
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
        max-width: 98%;
        width: 100%;
        color: #fff;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# --- CUSTOM CSS ---
def load_custom_css():
    st.markdown("""
    <style>
    /* ... existing CSS ... */
    /* Sidebar enhancements */
    .custom-sidebar {
        background: linear-gradient(180deg, #232526 0%, #414345 100%);
        border-radius: 0 24px 24px 0;
        padding: 1.5rem 1rem 2rem 1rem;
        min-height: 100vh;
        box-shadow: 2px 0 16px rgba(44,62,80,0.08);
        font-family: 'Inter', sans-serif;
        position: relative;
    }
    .custom-sidebar .sidebar-logo {
        display: block;
        margin: 0 auto 1.2rem auto;
        width: 64px;
        height: 64px;
        border-radius: 50%;
        background: linear-gradient(135deg, #667eea, #764ba2);
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 2.5rem;
        color: #fff;
        box-shadow: 0 4px 16px rgba(52, 152, 219, 0.15);
    }
    .custom-sidebar .sidebar-title {
        color: #ecf0f1;
        text-align: center;
        font-size: 1.5rem;
        font-weight: 700;
        margin-bottom: 0.2rem;
        letter-spacing: 1px;
    }
    .custom-sidebar .sidebar-subtitle {
        color: #bdc3c7;
        text-align: center;
        font-size: 1rem;
        margin-bottom: 1.2rem;
    }
    .custom-sidebar .sidebar-nav label {
        display: block;
        padding: 0.8rem 1.2rem;
        border-radius: 12px;
        margin-bottom: 0.5rem;
        color: #ecf0f1;
        font-weight: 500;
        transition: background 0.2s, color 0.2s;
        cursor: pointer;
    }
    .custom-sidebar .sidebar-nav label[data-checked="true"] {
        background: linear-gradient(135deg, #3498db, #2980b9);
        color: #fff;
        box-shadow: 0 2px 8px rgba(52, 152, 219, 0.15);
    }
    .custom-sidebar .sidebar-nav label:hover {
        background: rgba(255,255,255,0.08);
        color: #fff;
    }
    .custom-sidebar .custom-divider {
        margin: 2rem 0 1rem 0;
    }
    .custom-sidebar .sidebar-footer {
        color: #95a5a6;
        font-size: 0.85rem;
        text-align: center;
        margin-top: 2rem;
    }
    /* Chat history toggle button */
    .sidebar-btn {
        background: linear-gradient(135deg, #3498db, #2980b9);
        color: #fff !important;
        border: none;
        border-radius: 10px;
        font-size: 1.1rem;
        font-weight: 600;
        padding: 0.7rem 1.5rem;
        margin-bottom: 0.5rem;
        box-shadow: 0 2px 8px rgba(52, 152, 219, 0.12);
        transition: background 0.2s, transform 0.2s;
        width: 100%;
        cursor: pointer;
    }
    .sidebar-btn:hover {
        background: linear-gradient(135deg, #2980b9, #3498db);
        transform: translateY(-2px) scale(1.03);
    }
    /* File details card wider */
    .file-details-card {
        min-width: 260px;
        max-width: 420px;
        width: 100%;
        background: linear-gradient(135deg, rgba(255,255,255,0.97), rgba(255,255,255,0.85));
        border-radius: 16px;
        box-shadow: 0 4px 16px rgba(52, 152, 219, 0.08);
        padding: 1.5rem 1.2rem;
        margin: 0 auto 1rem auto;
        border: 1px solid rgba(52, 152, 219, 0.08);
    }
    </style>
    """, unsafe_allow_html=True)

# --- SIDEBAR LOGO & NAVIGATION ---
with st.sidebar:
    st.markdown('<div class="custom-sidebar">', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-logo">⚖️</div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-title">Legal AI</div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-subtitle">Professional Suite</div>', unsafe_allow_html=True)
    pages = [
        "🏠 Home",
        "📝 Document Drafting", 
        "🔍 Legal Clarification",
        "📄 Document-Based QA"
    ]
    page = st.radio("", ["🏠 Home", "📝 Document Drafting", "🔍 Legal Clarification", "📄 Document-Based QA"], index=0, key="sidebar_nav", label_visibility="collapsed")
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-footer">Unified Legal AI Suite © 2024<br>Built with ❤️ for Legal Professionals</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# --- SESSION STATE INIT ---
def init_session():
    if 'drafting_agent' not in st.session_state:
        st.session_state.drafting_agent = LegalDocumentAgent()
    if 'drafting_session_id' not in st.session_state:
        st.session_state.drafting_session_id = str(uuid.uuid4())
    if 'drafting_document' not in st.session_state:
        st.session_state.drafting_document = None
    if 'clarification_search' not in st.session_state:
        st.session_state.clarification_search = LegalSearchGraph(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            gemini_api_key=os.getenv("GEMINI_API_KEY")
        )
    if 'clarification_summarizer' not in st.session_state:
        st.session_state.clarification_summarizer = LegalSummarizer(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            gemini_api_key=os.getenv("GEMINI_API_KEY")
        )
    if 'clarification_history' not in st.session_state:
        st.session_state.clarification_history = []
    if 'docqa_rag' not in st.session_state:
        st.session_state.docqa_rag = DocumentQARAG(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            google_api_key=os.getenv("GEMINI_API_KEY")
        )
    if 'docqa_processed' not in st.session_state:
        st.session_state.docqa_processed = False
    if 'docqa_current_document' not in st.session_state:
        st.session_state.docqa_current_document = None
    if 'docqa_chat_history' not in st.session_state:
        st.session_state.docqa_chat_history = []

init_session()

# --- PAGE 1: HOME ---
if page == "🏠 Home":
    st.markdown('<h1 class="main-header">⚖️ Unified Legal AI Suite</h1>', unsafe_allow_html=True)
    st.markdown("""
    <div style='text-align:center; color:#fff; margin-bottom:3rem; font-size:1.2rem; font-weight:500;'>
        Draft legal documents, get legal clarifications, and ask questions about your own documents—all in one place.
    </div>
    """, unsafe_allow_html=True)
    # Three feature cards in a row
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown('''
        <div class="feature-card">
            <div style="text-align: center; margin-bottom: 1rem;">
                <div style="font-size: 3rem; margin-bottom: 0.5rem;">📝</div>
                <h3 style="margin: 0;">Document Drafting</h3>
            </div>
            <p style="line-height: 1.6;">Generate well-structured legal documents through conversational AI. From NDAs to lease agreements, our AI collects all required information and produces professional documents.</p>
            <div style="margin-top: 1rem;">
                <span style="background: rgba(255,255,255,0.15); color: #fff; padding: 4px 8px; border-radius: 12px; font-size: 0.8rem; font-weight: 500;">✨ Interactive</span>
                <span style="background: rgba(255,255,255,0.15); color: #fff; padding: 4px 8px; border-radius: 12px; font-size: 0.8rem; font-weight: 500; margin-left: 0.5rem;">📄 Professional</span>
            </div>
        </div>
        ''', unsafe_allow_html=True)
    with col2:
        st.markdown('''
        <div class="feature-card">
            <div style="text-align: center; margin-bottom: 1rem;">
                <div style="font-size: 3rem; margin-bottom: 0.5rem;">🔍</div>
                <h3 style="margin: 0;">Legal Clarification</h3>
            </div>
            <p style="line-height: 1.6;">Ask law-related questions and receive real-time, source-cited answers. Our system searches trusted legal databases and provides comprehensive summaries.</p>
            <div style="margin-top: 1rem;">
                <span style="background: rgba(255,255,255,0.15); color: #fff; padding: 4px 8px; border-radius: 12px; font-size: 0.8rem; font-weight: 500;">🔗 Source-Cited</span>
                <span style="background: rgba(255,255,255,0.15); color: #fff; padding: 4px 8px; border-radius: 12px; font-size: 0.8rem; font-weight: 500; margin-left: 0.5rem;">⚡ Real-time</span>
            </div>
        </div>
        ''', unsafe_allow_html=True)
    with col3:
        st.markdown('''
        <div class="feature-card">
            <div style="text-align: center; margin-bottom: 1rem;">
                <div style="font-size: 3rem; margin-bottom: 0.5rem;">📄</div>
                <h3 style="margin: 0;">Document-Based QA</h3>
            </div>
            <p style="line-height: 1.6;">Upload your legal documents and ask context-specific questions. Our AI analyzes your documents and provides accurate, grounded answers.</p>
            <div style="margin-top: 1rem;">
                <span style="background: rgba(255,255,255,0.15); color: #fff; padding: 4px 8px; border-radius: 12px; font-size: 0.8rem; font-weight: 500;">🎯 Context-Aware</span>
                <span style="background: rgba(255,255,255,0.15); color: #fff; padding: 4px 8px; border-radius: 12px; font-size: 0.8rem; font-weight: 500; margin-left: 0.5rem;">🔒 Secure</span>
            </div>
        </div>
        ''', unsafe_allow_html=True)
    st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
    # Technical Excellence card below
    st.markdown("""
    <div class="feature-card" style="margin-top: 2rem;">
        <h3 style="margin-bottom: 1rem;">🚀 Technical Excellence</h3>
        <div style="display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 1rem;">
            <div>
                <h4 style="margin-bottom: 0.5rem;">🧠 Advanced AI</h4>
                <p style="margin: 0;">Powered by state-of-the-art LLMs with robust fallback mechanisms and error handling.</p>
            </div>
            <div>
                <h4 style="margin-bottom: 0.5rem;">🔍 Vector Search</h4>
                <p style="margin: 0;">Intelligent document chunking and embedding for precise context retrieval.</p>
            </div>
            <div>
                <h4 style="margin-bottom: 0.5rem;">💾 Memory System</h4>
                <p style="margin: 0;">Contextual conversation memory for seamless multi-turn interactions.</p>
            </div>
            <div>
                <h4 style="margin-bottom: 0.5rem;">🔐 Privacy First</h4>
                <p style="margin: 0;">No sensitive data storage, educational use only, privacy-focused design.</p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# --- PAGE 2: DOCUMENT DRAFTING ---
elif page == "📝 Document Drafting":
    st.markdown('<h2 style="color:#1f4e79;">📝 Conversational Legal Document Drafting</h2>', unsafe_allow_html=True)
    st.info("Interactively draft NDAs, contracts, or lease agreements. The AI will ask for missing details and generate a complete document. Your chat history will be shown below.")
    def reset_drafting():
        st.session_state.drafting_session_id = str(uuid.uuid4())
        st.session_state.drafting_document = None
        st.session_state.drafting_chat = []
        if 'drafting_state_dict' in st.session_state:
            del st.session_state['drafting_state_dict']
        if 'drafting_question_index' in st.session_state:
            del st.session_state['drafting_question_index']
    st.button("Start New Drafting", on_click=reset_drafting)
    agent = st.session_state.drafting_agent
    if 'drafting_chat' not in st.session_state:
        st.session_state.drafting_chat = []
    # Only show input if no document is generated yet
    if not st.session_state.drafting_document:
        if 'drafting_state_dict' not in st.session_state:
            user_input = st.text_input("Describe the document you want to draft (e.g., 'Draft an NDA between Alice and Bob'):", key="drafting_input")
            if user_input:
                state = AgentState(session_id=st.session_state.drafting_session_id, user_input=user_input)
                state_dict = state.model_dump()
                state_dict = agent.identify_document_type(state_dict)
                st.session_state.drafting_state_dict = state_dict
                st.session_state.drafting_answers = {}
                st.session_state.drafting_question_index = 0
                st.session_state.drafting_chat.append({"role": "user", "content": user_input})
                st.rerun()
        else:
            state_dict = st.session_state.drafting_state_dict
            # Sequentially ask questions
            while not state_dict.get("is_complete", False):
                state_dict = agent.ask_question(state_dict)
                question = state_dict.get("current_question", "")
                if question:
                    q_idx = st.session_state.get('drafting_question_index', 0)
                    key = f"drafting_q_{st.session_state.drafting_session_id}_{q_idx}"
                    answer = st.text_input(f"{question}", key=key)
                    if answer:
                        state_dict["user_input"] = answer
                        state_dict = agent.process_answer(state_dict)
                        st.session_state.drafting_question_index = q_idx + 1
                        st.session_state.drafting_state_dict = state_dict
                        st.session_state.drafting_chat.append({"role": "ai", "content": question})
                        st.session_state.drafting_chat.append({"role": "user", "content": answer})
                        st.rerun()
                    else:
                        break
            # Generate document if complete
            if state_dict.get("is_complete", False):
                state_dict = agent.generate_document(state_dict)
                document = state_dict.get("final_document", "[No document generated]")
                st.session_state.drafting_document = document
                st.session_state.drafting_chat.append({"role": "ai", "content": "---\n**Generated Legal Document**\n" + document})
                del st.session_state['drafting_state_dict']
                del st.session_state['drafting_question_index']
                st.success("Document generated! See below.")
    # Show chat history with toggle
    if 'show_drafting_chat' not in st.session_state:
        st.session_state.show_drafting_chat = True
    if st.session_state.drafting_chat:
        st.markdown('<div class="custom-divider"></div>', unsafe_allow_html=True)
        col1, col2 = st.columns([1, 8])
        with col1:
            toggle_label = '👁️ Hide Chat' if st.session_state.show_drafting_chat else '👁️ Show Chat'
            if st.button(toggle_label, key='toggle_drafting_chat', help='Toggle chat history'):
                st.session_state.show_drafting_chat = not st.session_state.show_drafting_chat
        with col2:
            if st.session_state.show_drafting_chat:
                st.markdown("---")
                st.subheader("Chat History")
                # Only show Q&A, not the generated document
                for msg in st.session_state.drafting_chat:
                    if msg["role"] == "user":
                        st.markdown(f"**You:** {msg['content']}")
                    elif msg["role"] == "ai" and not msg["content"].startswith("---\n**Generated Legal Document**"):
                        st.markdown(f"**AI:** {msg['content']}")
    # Show the generated document and download button
    if st.session_state.drafting_document:
        st.markdown("---")
        # Format the document for markdown preview
        def format_legal_doc_to_markdown(doc_text):
            lines = doc_text.splitlines()
            md_lines = []
            for line in lines:
                line = line.strip()
                if not line:
                    md_lines.append("")
                    continue
                # Bold all-caps section headers and key phrases
                if line.isupper() and len(line) > 3:
                    md_lines.append(f"**{line.title()}**")
                elif line.startswith("LANDLORD:") or line.startswith("TENANT:") or line.startswith("PROPERTY:") or line.startswith("LEASE TERM:") or line.startswith("MONTHLY RENT:") or line.startswith("SECURITY DEPOSIT:") or line.startswith("USE OF PREMISES:") or line.startswith("CONDITION OF PREMISES:") or line.startswith("REPAIRS AND MAINTENANCE:") or line.startswith("UTILITIES:") or line.startswith("INSPECTIONS:") or line.startswith("DEFAULT:") or line.startswith("NOTICES:") or line.startswith("GOVERNING LAW:") or line.startswith("ENTIRE AGREEMENT:") or line.startswith("AMENDMENTS:") or line.startswith("BINDING EFFECT:") or line.startswith("ACKNOWLEDGMENT:") or line.startswith("SIGNATURES:"):
                    parts = line.split(":", 1)
                    if len(parts) == 2:
                        md_lines.append(f"**{parts[0]}:**{parts[1]}")
                    else:
                        md_lines.append(f"**{line}**")
                else:
                    md_lines.append(line)
            return "\n".join(md_lines)
        formatted_doc = format_legal_doc_to_markdown(st.session_state.drafting_document)
        st.markdown(f"{formatted_doc}")

        # Download as Word (.docx)
        def generate_docx(doc_text):
            doc = Document()
            lines = doc_text.splitlines()
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                # Bold all-caps section headers and key phrases
                if line.isupper() and len(line) > 3:
                    p = doc.add_paragraph()
                    run = p.add_run(line.title())
                    run.bold = True
                elif line.startswith("LANDLORD:") or line.startswith("TENANT:") or line.startswith("PROPERTY:") or line.startswith("LEASE TERM:") or line.startswith("MONTHLY RENT:") or line.startswith("SECURITY DEPOSIT:") or line.startswith("USE OF PREMISES:") or line.startswith("CONDITION OF PREMISES:") or line.startswith("REPAIRS AND MAINTENANCE:") or line.startswith("UTILITIES:") or line.startswith("INSPECTIONS:") or line.startswith("DEFAULT:") or line.startswith("NOTICES:") or line.startswith("GOVERNING LAW:") or line.startswith("ENTIRE AGREEMENT:") or line.startswith("AMENDMENTS:") or line.startswith("BINDING EFFECT:") or line.startswith("ACKNOWLEDGMENT:") or line.startswith("SIGNATURES:"):
                    parts = line.split(":", 1)
                    if len(parts) == 2:
                        p = doc.add_paragraph()
                        run = p.add_run(parts[0] + ":")
                        run.bold = True
                        p.add_run(parts[1])
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        run.bold = True
                else:
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf
        docx_buf = generate_docx(st.session_state.drafting_document)
        st.download_button(
            label="💾 Download Document (.docx)",
            data=docx_buf,
            file_name=f"legal_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- PAGE 3: LEGAL CLARIFICATION ---
elif page == "🔍 Legal Clarification":
    st.markdown('<h2 style="color:#1f4e79;">🔍 Legal Clarification</h2>', unsafe_allow_html=True)
    st.info("Ask a legal question. See sample questions below, extracted keywords, search results, and the LLM-generated summary.")
    sample_questions = [
        "What is the difference between void and voidable contracts in Canadian law?",
        "What are the requirements for adverse possession in Ontario?",
        "How is child custody determined in Canadian family courts?",
        "What constitutes negligence under Canadian tort law?"
    ]
    st.markdown("**Sample Questions:**")
    for q in sample_questions:
        st.markdown(f"- {q}")
    query = st.text_area(
        "Enter your legal question:",
        value=st.session_state.get("clarification_current_query", ""),
        height=80,
        placeholder="e.g., What is the difference between void and voidable contracts in Canadian law?",
        key="clarification_query_input"
    )
    col1, col2 = st.columns([1, 1])
    with col1:
        search_button = st.button("🔎 Search & Summarize", key="clarification_search_btn")
    with col2:
        clear_button = st.button("🗑️ Clear", key="clarification_clear_btn")
    if clear_button:
        st.session_state.clarification_current_query = ""
        st.rerun()
    if search_button and query.strip():
        st.session_state.clarification_current_query = query
        progress_bar = st.progress(0)
        status_text = st.empty()
        try:
            status_text.text("Searching legal databases...")
            progress_bar.progress(25)
            search_results = st.session_state.clarification_search.search_legal_query(query)
            if not search_results.get("success", False):
                st.error(f"Search failed: {search_results.get('error', 'Unknown error')}")
            else:
                status_text.text("Summarizing results...")
                progress_bar.progress(75)
                summary_results = st.session_state.clarification_summarizer.generate_comprehensive_summary(search_results)
                progress_bar.progress(100)
                status_text.text("Done!")
                time.sleep(0.5)
                progress_bar.empty()
                status_text.empty()
                st.markdown("---")
                st.markdown(f"**Your Question:** {query}")
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown("**Extracted Keywords:**")
                keywords = search_results.get("keywords", [])
                if keywords:
                    st.markdown(", ".join([f"`{kw}`" for kw in keywords]))
                else:
                    st.markdown("_No keywords extracted._")
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown("**Search Results:**")
                results = search_results.get("results", [])
                if results:
                    for i, result in enumerate(results, 1):
                        st.markdown(f"**{i}.** {result.get('content', '')[:300]}{'...' if len(result.get('content', ''))>300 else ''}")
                        st.caption(f"Source: {result.get('source', 'Unknown')}")
                else:
                    st.markdown("_No search results found._")
                st.markdown("<hr>", unsafe_allow_html=True)
                st.markdown("**LLM-Generated Summary:**")
                if summary_results.get("success", False):
                    st.success(summary_results.get('summary', ''))
                    if summary_results.get("citations"):
                        st.markdown("**Citations:**")
                        for i, citation in enumerate(summary_results["citations"], 1):
                            st.markdown(f"{i}. {citation}")
                else:
                    st.error(f"Summary failed: {summary_results.get('error', 'Unknown error')}")
        except Exception as e:
            st.error(f"Processing failed: {str(e)}")
            progress_bar.empty()
            status_text.empty()

# --- PAGE 4: DOCUMENT-BASED QA ---
elif page == "📄 Document-Based QA":
    st.markdown('<h2 style="color:#1f4e79;">📄 Document-Based QA</h2>', unsafe_allow_html=True)
    st.info("Upload a legal document (PDF, DOCX, or TXT) and ask context-specific questions.")
    col1, col2 = st.columns([2, 1])
    with col1:
        uploaded_file = st.file_uploader(
            "Choose a document",
            type=['pdf', 'docx', 'doc', 'txt'],
            help="Upload a PDF, DOCX, or TXT file"
        )
    with col2:
        if uploaded_file is not None:
            file_size_mb = uploaded_file.size / (1024 * 1024)
            st.markdown(f'''
            <div class="file-details-card">
                <h5 style="margin: 0 0 0.5rem 0; color: #2c3e50;">📄 File Details</h5>
                <p style="margin: 0; color: #555;"><strong>Name:</strong> {uploaded_file.name}</p>
                <p style="margin: 0; color: #555;"><strong>Size:</strong> {file_size_mb:.1f} MB</p>
                <p style="margin: 0; color: #555;"><strong>Type:</strong> {uploaded_file.type}</p>
            </div>
            ''', unsafe_allow_html=True)
    if uploaded_file is not None:
        if st.button("🚀 Process Document", key="docqa_process_btn"):
            with st.spinner("Processing document..."):
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                file_extension = uploaded_file.name.split('.')[-1].lower()
                try:
                    result = st.session_state.docqa_rag.process_document_and_query(
                        file_path=tmp_file_path,
                        file_type=file_extension,
                        query="What is this document about?"
                    )
                except Exception as e:
                    result = {'success': False, 'error': str(e)}
                os.unlink(tmp_file_path)
                if result['success']:
                    st.session_state.docqa_processed = True
                    st.session_state.docqa_current_document = uploaded_file.name
                    st.success(f"Document processed! You can now ask questions about: {uploaded_file.name}")
                    st.write("**Initial Analysis:**")
                    st.write(result['answer'])
                else:
                    st.session_state.docqa_processed = False
                    st.session_state.docqa_current_document = None
                    error_msg = result.get('error', 'Unknown error occurred')
                    if error_msg is None:
                        error_msg = 'Unknown error occurred'
                    if isinstance(error_msg, str):
                        if 'no healthy upstream' in error_msg or 'No LLM could be initialized' in error_msg:
                            st.error("All LLMs are currently unavailable. Please try again later or check your API keys.")
                        else:
                            st.error(f"Error processing document: {error_msg}")
                    else:
                        st.error("An unknown error occurred while processing the document.")
    if st.session_state.docqa_processed:
        st.markdown("---")
        st.subheader("💬 Ask Questions About Your Document")
        question = st.text_input(
            "Ask a question about your document:",
            placeholder="e.g., What are the main obligations of the parties?",
            key="docqa_question_input"
        )
        ask_button = st.button("🤔 Ask Question", key="docqa_ask_btn")
        clear_history = st.button("🗑️ Clear History", key="docqa_clear_btn")
        if clear_history:
            st.session_state.docqa_chat_history = []
            st.rerun()
        if ask_button and question:
            with st.spinner("Thinking..."):
                result = st.session_state.docqa_rag.query_existing_documents(question)
            if result['success']:
                timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
                st.session_state.docqa_chat_history.append((question, result['answer'], timestamp))
                st.success(result['answer'])
            else:
                st.error(f"Error: {result.get('error', 'Unknown error occurred')}")
        if st.session_state.docqa_chat_history:
            st.subheader("💭 Chat History")
            for i, (q, a, timestamp) in enumerate(reversed(st.session_state.docqa_chat_history)):
                with st.expander(f"Q{len(st.session_state.docqa_chat_history)-i}: {q[:50]}..." if len(q) > 50 else f"Q{len(st.session_state.docqa_chat_history)-i}: {q}", expanded=(i==0)):
                    st.write(f"**Question:** {q}")
                    st.write(f"**Answer:** {a}")
                    st.caption(f"Asked at: {timestamp}")

# --- FOOTER ---
st.markdown("""
<div class='footer'>
  <span>Made with <span style='color:#ffd600;'>&#9733;</span> by Legal AI Suite Team &copy; 2024</span>
</div>
""", unsafe_allow_html=True) 